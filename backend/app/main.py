import io
import re
import textwrap
import unicodedata
from pathlib import Path
from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Form, status
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import RedirectResponse, FileResponse, StreamingResponse
from fastapi.security import OAuth2PasswordRequestForm
from fastapi.staticfiles import StaticFiles
from sqlalchemy import inspect, text
from sqlalchemy.orm import Session
from pypdf import PdfReader
from docx import Document
from fpdf import FPDF
import requests
from bs4 import BeautifulSoup
from groq import RateLimitError

from app import models, schemas, auth
from app.database import engine, get_db
from app.agents.coordinator import AgentCoordinator

# Initialize DB Tables
models.Base.metadata.create_all(bind=engine)

def ensure_db_columns():
    inspector = inspect(engine)
    if "applications" in inspector.get_table_names():
        app_columns = [col["name"] for col in inspector.get_columns("applications")]
        if "jd_url" not in app_columns:
            with engine.begin() as conn:
                conn.execute(text("ALTER TABLE applications ADD COLUMN jd_url TEXT"))
        if "resume_sections" not in app_columns:
            with engine.begin() as conn:
                conn.execute(text("ALTER TABLE applications ADD COLUMN resume_sections TEXT"))
    if "drafts" in inspector.get_table_names():
        draft_columns = [col["name"] for col in inspector.get_columns("drafts")]
        if "ats_score" not in draft_columns:
            with engine.begin() as conn:
                conn.execute(text("ALTER TABLE drafts ADD COLUMN ats_score TEXT"))

ensure_db_columns()

app = FastAPI(title="Job Application Co-Pilot API")
static_dir = Path(__file__).resolve().parent / "static"
app.mount("/static", StaticFiles(directory=static_dir), name="static")

@app.get("/", include_in_schema=False)
def root():
    return FileResponse(static_dir / "index.html")

# Enable CORS for vanilla JS or local testing environments
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

coordinator = AgentCoordinator()

def parse_resume_sections(resume_text: str) -> str:
    headings = [
        "SUMMARY", "PROFILE", "EXPERIENCE", "WORK EXPERIENCE",
        "EDUCATION", "SKILLS", "PROJECTS", "CERTIFICATIONS",
        "CERTIFICATIONS & AWARDS", "ACHIEVEMENTS"
    ]
    sections = {"Summary": []}
    current = "Summary"

    for line in resume_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue
        normalized = re.sub(r"[^A-Z0-9 ]", "", stripped.upper()).strip()
        matched = False
        for heading in headings:
            if normalized.startswith(heading):
                current = heading.title()
                sections.setdefault(current, [])
                matched = True
                break
        if not matched:
            sections.setdefault(current, []).append(stripped)

    formatted = []
    for name, lines in sections.items():
        if lines:
            formatted.append(f"{name}:\n" + "\n".join(lines))
    return "\n\n".join(formatted)

def scrape_jd_text(url: str) -> str:
    try:
        response = requests.get(url, timeout=10, headers={"User-Agent": "Mozilla/5.0"})
        response.raise_for_status()
        soup = BeautifulSoup(response.text, "html.parser")
        text_blocks = [tag.get_text(" ", strip=True) for tag in soup.find_all(["h1", "h2", "h3", "h4", "p", "li", "dd", "dt"])]
        scraped_text = "\n\n".join([block for block in text_blocks if block])
        if not scraped_text.strip():
            raise ValueError("Could not extract job description text from the provided URL")
        return scraped_text
    except Exception as exc:
        raise HTTPException(status_code=400, detail=f"Failed to scrape JD URL: {exc}")

@app.post("/signup", response_model=schemas.Token)
def signup(user_data: schemas.UserCreate, db: Session = Depends(get_db)):
    db_user = db.query(models.User).filter(models.User.username == user_data.username).first()
    if db_user:
        raise HTTPException(status_code=400, detail="Username already registered")
    
    hashed_pwd = auth.get_password_hash(user_data.password)
    new_user = models.User(username=user_data.username, hashed_password=hashed_pwd)
    db.add(new_user)
    db.commit()
    db.refresh(new_user)
    
    access_token = auth.create_access_token(data={"sub": new_user.username})
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/token", response_model=schemas.Token)
def login(form_data: OAuth2PasswordRequestForm = Depends(), db: Session = Depends(get_db)):
    user = db.query(models.User).filter(models.User.username == form_data.username).first()
    if not user or not auth.verify_password(form_data.password, user.hashed_password):
        raise HTTPException(status_code=400, detail="Incorrect username or password")
    
    access_token = auth.create_access_token(data={"sub": user.username})
    return {"access_token": access_token, "token_type": "bearer"}

@app.post("/applications", status_code=status.HTTP_201_CREATED)
async def create_application(
    job_title: str = Form(...),
    company: str = Form(...),
    jd_text: str = Form(""),
    jd_url: str | None = Form(None),
    resume_file: UploadFile = File(...),
    current_user: models.User = Depends(auth.get_current_user),
    db: Session = Depends(get_db)
):
    try:
        pdf_content = await resume_file.read()
        pdf_reader = PdfReader(io.BytesIO(pdf_content))
        resume_text = ""
        for page in pdf_reader.pages:
            text = page.extract_text()
            if text:
                resume_text += text + "\n"
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"Failed to process PDF resume: {str(e)}")

    if not resume_text.strip():
        raise HTTPException(status_code=400, detail="Could not extract readable text from the uploaded PDF resume.")

    if not jd_text and jd_url:
        jd_text = scrape_jd_text(jd_url)
    if not jd_text:
        raise HTTPException(status_code=400, detail="Job description text or JD URL is required.")

    resume_sections = parse_resume_sections(resume_text)

    db_app = models.Application(
        job_title=job_title,
        company=company,
        jd_text=jd_text,
        jd_url=jd_url,
        original_resume_text=resume_text,
        resume_sections=resume_sections,
        user_id=current_user.id
    )
    db.add(db_app)
    db.commit()
    db.refresh(db_app)

    try:
        pipeline_outputs = coordinator.run_pipeline(resume_text, jd_text)
    except RateLimitError as e:
        raise HTTPException(
            status_code=429,
            detail=f"AI Pipeline Execution Error: Rate limit reached. {str(e)}"
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"AI Pipeline Execution Error: {str(e)}")

    db_draft = models.Draft(
        application_id=db_app.id,
        fit_analysis=pipeline_outputs["fit_analysis"],
        resume_rewrite=pipeline_outputs["resume_rewrite"],
        cover_letter=pipeline_outputs["cover_letter"],
        interview_qa=pipeline_outputs["interview_qa"],
        ats_score=pipeline_outputs["ats_score"]
    )
    db.add(db_draft)
    db.commit()

    return {
        "message": "Application kit created successfully",
        "application_id": db_app.id,
        "artifacts": pipeline_outputs
    }

@app.get("/applications", response_model=list[schemas.ApplicationResponse])
def list_applications(current_user: models.User = Depends(auth.get_current_user), db: Session = Depends(get_db)):
    return db.query(models.Application).filter(models.Application.user_id == current_user.id).all()

@app.get("/applications/{app_id}", response_model=schemas.ApplicationDetailResponse)
def get_application_detail(app_id: int, current_user: models.User = Depends(auth.get_current_user), db: Session = Depends(get_db)):
    application = db.query(models.Application).filter(models.Application.id == app_id, models.Application.user_id == current_user.id).first()
    if not application:
        raise HTTPException(status_code=404, detail="Application record not found")
    return application

@app.get("/applications/{app_id}/drafts", response_model=schemas.DraftResponse)
def get_application_drafts(app_id: int, current_user: models.User = Depends(auth.get_current_user), db: Session = Depends(get_db)):
    application = db.query(models.Application).filter(models.Application.id == app_id, models.Application.user_id == current_user.id).first()
    if not application:
        raise HTTPException(status_code=404, detail="Application record not found")
    if not application.drafts:
        raise HTTPException(status_code=404, detail="Drafts have not been generated for this application")
    return application.drafts

@app.put("/applications/{app_id}/status")
def update_application_status(
    app_id: int, 
    status_update: schemas.ApplicationStatusUpdate, 
    current_user: models.User = Depends(auth.get_current_user), 
    db: Session = Depends(get_db)
):
    application = db.query(models.Application).filter(models.Application.id == app_id, models.Application.user_id == current_user.id).first()
    if not application:
        raise HTTPException(status_code=404, detail="Application record not found")
    
    application.status = status_update.status
    db.commit()
    return {"message": "Status updated successfully", "new_status": application.status}

@app.put("/applications/{app_id}/regenerate")
def regenerate_application_section(
    app_id: int,
    request: schemas.RegenerateRequest,
    current_user: models.User = Depends(auth.get_current_user),
    db: Session = Depends(get_db)
):
    application = db.query(models.Application).filter(models.Application.id == app_id, models.Application.user_id == current_user.id).first()
    if not application or not application.drafts:
        raise HTTPException(status_code=404, detail="Application or drafts not found")

    valid_sections = {"fit_analysis", "resume_rewrite", "cover_letter", "interview_qa", "ats_score"}
    section = request.section.lower()
    if section not in valid_sections:
        raise HTTPException(status_code=400, detail=f"Supported sections: {', '.join(valid_sections)}")

    current_fit = application.drafts.fit_analysis
    current_resume = application.drafts.resume_rewrite
    try:
        updated_output = coordinator.regenerate_section(
            section,
            application.original_resume_text,
            application.jd_text,
            fit_analysis=current_fit,
            resume_rewrite=current_resume,
        )
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Regeneration error: {exc}")

    setattr(application.drafts, section, updated_output)
    db.commit()
    return {"message": f"Regenerated {section} successfully", section: updated_output}

@app.post("/applications/{app_id}/ats-score")
def compute_ats_score(app_id: int, current_user: models.User = Depends(auth.get_current_user), db: Session = Depends(get_db)):
    application = db.query(models.Application).filter(models.Application.id == app_id, models.Application.user_id == current_user.id).first()
    if not application or not application.drafts or not application.drafts.resume_rewrite:
        raise HTTPException(status_code=404, detail="Application or resume rewrite not available for ATS scoring")

    try:
        updated_score = coordinator.score_ats(application.drafts.resume_rewrite, application.jd_text)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"ATS scoring error: {exc}")

    application.drafts.ats_score = updated_score
    db.commit()
    return {"message": "ATS score computed successfully", "ats_score": updated_score}

@app.get("/applications/{app_id}/download/cover-letter")
def download_cover_letter(app_id: int, current_user: models.User = Depends(auth.get_current_user), db: Session = Depends(get_db)):
    application = db.query(models.Application).filter(models.Application.id == app_id, models.Application.user_id == current_user.id).first()
    if not application or not application.drafts or not application.drafts.cover_letter:
        raise HTTPException(status_code=404, detail="Cover letter not available")

    document = Document()
    document.add_heading(f"Cover Letter for {application.job_title} at {application.company}", level=1)
    for paragraph in application.drafts.cover_letter.split("\n\n"):
        document.add_paragraph(paragraph.strip())

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    filename = f"cover_letter_{application.job_title.replace(' ', '_')}.docx"
    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={filename}"}
    )

@app.get("/applications/{app_id}/download/resume")
def download_resume(app_id: int, current_user: models.User = Depends(auth.get_current_user), db: Session = Depends(get_db)):
    application = db.query(models.Application).filter(models.Application.id == app_id, models.Application.user_id == current_user.id).first()
    if not application or not application.drafts or not application.drafts.resume_rewrite:
        raise HTTPException(status_code=404, detail="Resume rewrite not available")

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    line_width = pdf.w - 2 * pdf.l_margin

    def normalize_for_pdf(text: str) -> str:
        normalized = unicodedata.normalize("NFKD", text)
        normalized = normalized.replace("\t", "    ")
        return normalized.encode("latin-1", errors="ignore").decode("latin-1")

    for line in application.drafts.resume_rewrite.replace('\r', '').split('\n'):
        if not line.strip():
            pdf.ln(8)
            continue
        safe_line = normalize_for_pdf(line)
        try:
            pdf.multi_cell(line_width, 8, safe_line)
        except Exception:
            current = ''
            for word in safe_line.split(' '):
                candidate = f"{current} {word}".strip()
                if pdf.get_string_width(candidate) <= line_width:
                    current = candidate
                    continue
                if current:
                    pdf.multi_cell(line_width, 8, current)
                if pdf.get_string_width(word) <= line_width:
                    current = word
                    continue
                # Break extremely long word into smaller chunks.
                fragment = ''
                for char in word:
                    if pdf.get_string_width(fragment + char) > line_width:
                        pdf.multi_cell(line_width, 8, fragment)
                        fragment = char
                    else:
                        fragment += char
                if fragment:
                    pdf.multi_cell(line_width, 8, fragment)
                current = ''
            if current:
                pdf.multi_cell(line_width, 8, current)

    buffer = io.BytesIO()
    pdf_bytes = pdf.output(dest="S")
    if isinstance(pdf_bytes, str):
        pdf_bytes = pdf_bytes.encode("latin-1")
    buffer.write(pdf_bytes)
    buffer.seek(0)
    safe_title = re.sub(r"[^A-Za-z0-9_\-]", "_", application.job_title.strip().replace(' ', '_'))
    filename = f"resume_{safe_title}.pdf"
    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename=\"{filename}\""}
    )